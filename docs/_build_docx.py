# Builds USE_CASES.docx from USE_CASES.md (headings, tables, lists, inline formatting,
# embedded diagram). Tailored to this document's markdown subset; not a general converter.
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "USE_CASES.md")
OUT = os.path.join(HERE, "USE_CASES.docx")

doc = Document()
# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

INLINE = re.compile(r"\*\*(?P<b>.+?)\*\*|`(?P<c>[^`]+)`|\*(?P<i>.+?)\*")

def add_runs(paragraph, text):
    """Add text to a paragraph, honoring **bold**, *italic*, and `code`."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("b") is not None:
            paragraph.add_run(m.group("b")).bold = True
        elif m.group("c") is not None:
            r = paragraph.add_run(m.group("c"))
            r.font.name = "Consolas"; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif m.group("i") is not None:
            paragraph.add_run(m.group("i")).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

lines = open(MD, encoding="utf-8").read().splitlines()
i = 0
n = len(lines)
while i < n:
    line = lines[i]
    stripped = line.strip()

    # Fenced code blocks
    if stripped.startswith("```"):
        lang = stripped[3:].strip().lower()
        block = []
        i += 1
        while i < n and not lines[i].strip().startswith("```"):
            block.append(lines[i]); i += 1
        i += 1  # skip closing fence
        if lang == "plantuml":
            continue  # represented by the embedded diagram image
        p = doc.add_paragraph()
        r = p.add_run("\n".join(block))
        r.font.name = "Consolas"; r.font.size = Pt(9)
        continue

    # Embedded image
    mimg = re.match(r"!\[(?P<alt>.*?)\]\((?P<src>.+?)\)", stripped)
    if mimg:
        src = os.path.join(HERE, mimg.group("src"))
        if os.path.exists(src):
            doc.add_picture(src, width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Tables (consecutive lines starting with |)
    if stripped.startswith("|"):
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells); i += 1
        # rows[1] is the |---| separator
        header = rows[0]
        body = rows[2:] if len(rows) > 1 else []
        t = doc.add_table(rows=1, cols=len(header))
        t.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            cell = t.rows[0].cells[j]
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(h); run.bold = True
        for r in body:
            cells = t.add_row().cells
            for j in range(len(header)):
                txt = r[j] if j < len(r) else ""
                cells[j].paragraphs[0].text = ""
                add_runs(cells[j].paragraphs[0], txt)
        doc.add_paragraph()
        continue

    # Headings
    if stripped.startswith("#"):
        m = re.match(r"(#+)\s+(.*)", stripped)
        level = len(m.group(1)) - 1  # '#' -> Title(0)
        doc.add_heading(m.group(2), level=min(level, 4))
        i += 1
        continue

    # Horizontal rule
    if stripped == "---":
        i += 1
        continue

    # Blockquote
    if stripped.startswith(">"):
        p = doc.add_paragraph()
        p.style = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
        add_runs(p, stripped.lstrip("> ").strip())
        i += 1
        continue

    # Lists (bullets / numbered, with 2-space nesting)
    indent = len(line) - len(line.lstrip(" "))
    mnum = re.match(r"\d+\.\s+(.*)", stripped)
    if stripped.startswith("- "):
        style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
        p = doc.add_paragraph(style=style_name)
        add_runs(p, stripped[2:])
        i += 1
        continue
    if mnum:
        style_name = "List Number 2" if indent >= 2 else "List Number"
        p = doc.add_paragraph(style=style_name)
        add_runs(p, mnum.group(1))
        i += 1
        continue

    # Blank or plain paragraph
    if stripped == "":
        i += 1
        continue
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

doc.save(OUT)
print("WROTE", OUT, os.path.getsize(OUT), "bytes")
